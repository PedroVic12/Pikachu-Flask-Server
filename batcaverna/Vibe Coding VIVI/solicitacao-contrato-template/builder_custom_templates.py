import docx
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL as WD_CELL_VERTICAL_ALIGN

def set_cell_background(cell, fill):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_dropdown(paragraph, options, default=None):
    if default is None: default = options[0]
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    dd = OxmlElement('w:dropDownList')
    for opt in options:
        listItem = OxmlElement('w:listItem')
        listItem.set(qn('w:displayText'), opt)
        listItem.set(qn('w:value'), opt)
        dd.append(listItem)
    sdtPr.append(dd)
    sdt.append(sdtPr)
    sdtContent = OxmlElement('w:sdtContent')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = default
    r.append(t)
    sdtContent.append(r)
    sdt.append(sdtContent)
    paragraph._p.append(sdt)

class ContractTemplateBuilder:
    def __init__(self, base_template_path):
        self.doc = docx.Document(base_template_path)
        self._clear_body()

    def _clear_body(self):
        for table in self.doc.tables:
            table._element.getparent().remove(table._element)
        for para in self.doc.paragraphs:
            para.text = ""

    def add_title(self, text):
        title = self.doc.add_paragraph(text)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run() # Title run was cleared, need to add new one
        run.text = text
        run.bold = True
        run.font.size = Pt(14)
        self.doc.add_paragraph()

    def create_form_table(self, config):
        col_widths = [15, 35, 15, 30, 15, 20, 15, 20, 10, 10]
        table = self.doc.add_table(rows=0, cols=10)
        table.style = 'Table Grid'
        table.autofit = False
        
        for i, width in enumerate(col_widths):
            table.columns[i].width = Mm(width)

        def add_row(height_mm=8):
            row = table.add_row()
            row.height = Mm(height_mm)
            return row

        def setup_cell(row, col, text, bold=False, bg=None, span_to_col=None, dropdown_options=None):
            cell = row.cells[col]
            if span_to_col is not None:
                cell.merge(row.cells[span_to_col])
            
            p = cell.paragraphs[0]
            if text:
                run = p.add_run(text)
                run.bold = bold
                run.font.size = Pt(8)
            if dropdown_options:
                add_dropdown(p, dropdown_options)
            if bg:
                set_cell_background(cell, bg)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGN.CENTER
            return cell

        # Dinamic Sections based on config
        # Row 1: Basic Info
        r = add_row()
        setup_cell(r, 0, "Solicitante:", bold=True)
        setup_cell(r, 1, "", span_to_col=5)
        setup_cell(r, 6, "Data Requisição:", bold=True)
        setup_cell(r, 7, "", span_to_col=9)

        r = add_row()
        setup_cell(r, 0, "Diretoria:", bold=True)
        setup_cell(r, 1, "")
        setup_cell(r, 2, "Ger. Exec:", bold=True)
        setup_cell(r, 3, "")
        setup_cell(r, 4, "Localidade:", bold=True)
        setup_cell(r, 5, "", span_to_col=9, dropdown_options=config.get('localidades', ['Brasília - DF']))

        # Row 3: Cargo & Motivo
        r = add_row()
        setup_cell(r, 0, "Cargo:", bold=True)
        setup_cell(r, 1, "", span_to_col=9)

        r = add_row()
        setup_cell(r, 0, "Motivo:", bold=True)
        setup_cell(r, 1, "", span_to_col=9, dropdown_options=config.get('motivos', ['Efetivo']))

        # Section GP
        r = add_row()
        setup_cell(r, 0, "Preenchimento pela GP", bold=True, bg="D9D9D9", span_to_col=9)
        
        r = add_row(height_mm=20)
        setup_cell(r, 0, "Comentários/Análise:", bold=True)
        setup_cell(r, 1, "", span_to_col=9)

        # Custom Fields for specific contract types
        if config.get('type') == 'TEMPORARIO':
            r = add_row()
            setup_cell(r, 0, "Prazo Contrato:", bold=True)
            setup_cell(r, 1, "", span_to_col=9)

    def add_approval_flow(self, flow_data):
        self.doc.add_paragraph()
        table = self.doc.add_table(rows=len(flow_data), cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r, row_data in enumerate(flow_data):
            for c, text in enumerate(row_data):
                cell = table.cell(r, c)
                cell.text = text
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.font.size = Pt(8)
                if r < 2:
                    run.bold = True
                    set_cell_background(cell, "D9D9D9")

    def save(self, path):
        self.doc.save(path)
        print(f"Template customizado salvo em: {path}")

if __name__ == "__main__":
    # Example: Creating a "Temporário" Contract Template
    config_temp = {
        'type': 'TEMPORARIO',
        'motivos': ['Substituição Licença', 'Aumento Sazonal', 'Projeto Específico'],
        'localidades': ['Brasília - DF', 'Rio de Janeiro - RJ', 'São Paulo - SP']
    }
    
    builder = ContractTemplateBuilder('Solicitação de Contratação_tamplet.docx')
    builder.add_title("SOLICITAÇÃO DE CONTRATAÇÃO - TEMPORÁRIO")
    builder.create_form_table(config_temp)
    
    flow = [
        ["PARA APROVAÇÃO", "PARA APROVAÇÃO"],
        ["GERÊNCIA", "DIRETORIA"],
        ["Gestor Imediato", "Diretor de Área"],
        ["Gerente Executivo", "RH"]
    ]
    builder.add_approval_flow(flow)
    
    builder.save("Template_Contrato_Temporario.docx")
