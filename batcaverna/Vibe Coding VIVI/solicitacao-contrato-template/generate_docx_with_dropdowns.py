import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_dropdown(paragraph, label_text, options):
    """
    Inserts a dropdown (Structured Document Tag) into a paragraph.
    Note: Standard python-docx doesn't support SDTs directly, so we use oxml.
    """
    run = paragraph.add_run(label_text + " ")
    
    # Create the SDT (Structured Document Tag) element
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    
    # Dropdown ID and alias
    alias = OxmlElement('w:alias')
    alias.set(qn('w:val'), label_text.strip(':'))
    sdtPr.append(alias)
    
    # Dropdown type
    dd = OxmlElement('w:dropDownList')
    for opt in options:
        listItem = OxmlElement('w:listItem')
        listItem.set(qn('w:displayText'), opt)
        listItem.set(qn('w:value'), opt)
        dd.append(listItem)
    sdtPr.append(dd)
    
    sdt.append(sdtPr)
    
    # SDT Content
    sdtContent = OxmlElement('w:sdtContent')
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = options[0] # Default value
    r.append(t)
    p.append(r)
    # We append a run inside the sdtContent instead of a whole paragraph if possible
    # but for simplicity in complex tables, a run is better.
    # Actually, SDT in Word usually wraps a run.
    
    sdtContent_run = OxmlElement('w:r')
    sdtContent_t = OxmlElement('w:t')
    sdtContent_t.text = options[0]
    sdtContent_run.append(sdtContent_t)
    sdtContent.append(sdtContent_run)
    
    sdt.append(sdtContent)
    
    paragraph._p.append(sdt)

def modify_docx():
    doc = docx.Document('Solicitação de Contratação_tamplet.docx')
    
    # 1. Add Dropdown for Localidade
    # Finding the cell that contains "UA:" or similar to place Localidade
    found_ua = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "UA:" in cell.text:
                    # Append "Localidade:" and Dropdown
                    p = cell.paragraphs[0]
                    p.add_run(" | Localidade: ")
                    localidades = ['Brasília - DF', 'Rio de Janeiro - RJ', 'São Paulo - SP', 'Belo Horizonte - MG']
                    add_dropdown(p, "", localidades)
                    found_ua = True
                    break
            if found_ua: break
        if found_ua: break

    # 2. Add Dropdown for Motivo
    # The Motivo row in the template has multiple options as text. 
    # We will replace that with a dropdown.
    found_motivo = False
    for table in doc.tables:
        for row in table.rows:
            if "Motivo:" in row.cells[0].text:
                # Clear the cells that follow "Motivo:" and put a dropdown in the first empty cell
                # Based on previous inspection, Motivo is in row 6 (index)
                cell_with_options = row.cells[1]
                cell_with_options.text = "" # Clear current text
                p = cell_with_options.paragraphs[0]
                motivos = ['Aumento de quadro', 'Efetivo', 'Trainee (prazo determinado)', 'Efetivo (prazo determinado)']
                add_dropdown(p, "", motivos)
                
                # Clear other cells in that row that might contain the static options
                for i in range(2, len(row.cells)):
                    row.cells[i].text = ""
                    
                found_motivo = True
                break
        if found_motivo: break

    output_path = 'Solicitacao_Contratacao_Com_Dropdowns.docx'
    doc.save(output_path)
    print(f"Documento salvo com sucesso: {output_path}")

if __name__ == "__main__":
    modify_docx()
