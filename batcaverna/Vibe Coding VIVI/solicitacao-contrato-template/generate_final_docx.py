import docx
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL as WD_CELL_VERTICAL_ALIGN

def set_cell_background(cell, fill, color=None, val=None):
    """
    Sets the background color of a cell.
    """
    shading_elm_1 = OxmlElement('w:shd')
    shading_elm_1.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm_1)

def set_cell_border(cell, **kwargs):
    """
    Sets cell borders.
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#FF0000"}, ...)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            edge_elm = OxmlElement(f'w:{edge}')
            for key, val in edge_data.items():
                edge_elm.set(qn(f'w:{key}'), str(val))
            tcBorders.append(edge_elm)
    
    tcPr.append(tcBorders)

def add_dropdown(paragraph, options, default=None):
    """
    Inserts a Word Dropdown (SDT) into a paragraph.
    """
    if default is None:
        default = options[0]
        
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    
    # Dropdown type
    dd = OxmlElement('w:dropDownList')
    for opt in options:
        listItem = OxmlElement('w:listItem')
        listItem.set(qn('w:displayText'), opt)
        listItem.set(qn('w:value'), opt)
        dd.append(listItem)
    sdtPr.append(dd)
    sdt.append(sdtPr)
    
    sdtContent = OxmlElement('w:sdtContent')
    r = OxmlElement('w:r')
    # Style the text inside dropdown to be visible
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF') # Blue to mimic PDF fields
    rPr.append(color)
    r.append(rPr)
    
    t = OxmlElement('w:t')
    t.text = default
    r.append(t)
    sdtContent.append(r)
    sdt.append(sdtContent)
    
    paragraph._p.append(sdt)

def create_final_docx():
    # Load original template to keep header/sections
    doc = docx.Document('Solicitação de Contratação_tamplet.docx')
    
    # Clear the body paragraphs and tables but keep the first few for spacing if needed
    # Actually, let's just delete everything after the first few empty paragraphs
    for table in doc.tables:
        table._element.getparent().remove(table._element)
    
    for para in doc.paragraphs:
        if para.text.strip() == "FLUXO DE APROVAÇÃO POR CARGO" or para.text.strip() == "":
            para.text = "" # Clear it
        else:
            # If it's something else we don't recognize as header, clear it
            # But let's be safe and only clear the body
            para.text = ""

    # Title
    title = doc.add_paragraph("FLUXO DE APROVAÇÃO POR CARGO")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph() # Spacer

    # Main Table
    # 17 rows, 10 columns as in PDF script
    col_widths = [15, 35, 15, 30, 15, 20, 15, 20, 10, 10]
    table = doc.add_table(rows=17, cols=10)
    table.style = 'Table Grid'
    table.autofit = False
    
    for i, width in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = Mm(width)

    # Helper for cell content
    def setup_cell(row, col, text, bold=False, bg=None, span_to_col=None):
        cell = table.cell(row, col)
        if span_to_col is not None:
            cell.merge(table.cell(row, span_to_col))
        
        p = cell.paragraphs[0]
        p.margin_top = Pt(2)
        p.margin_bottom = Pt(2)
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(8)
        if bg:
            set_cell_background(cell, bg)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGN.CENTER
        return cell

    # Row 0: Solicitante & Data
    setup_cell(0, 0, "Solicitante:", bold=True)
    setup_cell(0, 1, "", span_to_col=5)
    setup_cell(0, 6, "Data da Requisição:", bold=True)
    setup_cell(0, 7, "", span_to_col=9)
    
    # Row 1: Diretoria, Ger. Exec, Ger, Localidade
    setup_cell(1, 0, "Diretoria:", bold=True)
    setup_cell(1, 1, "")
    setup_cell(1, 2, "Ger. Exec:", bold=True)
    setup_cell(1, 3, "")
    setup_cell(1, 4, "Gerência:", bold=True)
    setup_cell(1, 5, "")
    setup_cell(1, 6, "Localidade:", bold=True)
    loc_cell = setup_cell(1, 7, "", span_to_col=9)
    add_dropdown(loc_cell.paragraphs[0], ['Brasília - DF', 'Rio de Janeiro - RJ', 'São Paulo - SP', 'Belo Horizonte - MG'])
    
    # Row 2: Cargo
    setup_cell(2, 0, "Cargo:", bold=True)
    setup_cell(2, 1, "", span_to_col=9)
    
    # Row 3: Motivo
    setup_cell(3, 0, "Motivo:", bold=True)
    mot_cell = setup_cell(3, 1, "", span_to_col=9)
    add_dropdown(mot_cell.paragraphs[0], ['Aumento de quadro', 'Efetivo', 'Trainee (prazo determinado)', 'Efetivo (prazo determinado)'])
    
    # Row 4: Em substituição
    setup_cell(4, 0, "Em substituição a:", bold=True)
    setup_cell(4, 1, "", span_to_col=9)
    
    # Row 5: Observação
    setup_cell(5, 0, "Observação:", bold=True)
    setup_cell(5, 1, "", span_to_col=9)
    table.rows[5].height = Mm(20)
    
    # Row 6: Header GP
    setup_cell(6, 0, "Preenchimento pela GP", bold=True, bg="D9D9D9", span_to_col=9)
    
    # Row 7: Análise Remuneração
    setup_cell(7, 0, "Análise da Remuneração", bold=True, span_to_col=9)
    
    # Row 8: Faixa salarial
    setup_cell(8, 0, "Faixa salarial:", bold=True)
    setup_cell(8, 1, "Inicial:")
    setup_cell(8, 2, "")
    setup_cell(8, 3, "Mediana:")
    setup_cell(8, 4, "")
    setup_cell(8, 5, "Final:")
    setup_cell(8, 6, "")
    setup_cell(8, 7, "Salário Admissão:")
    setup_cell(8, 8, "", span_to_col=9)
    
    # Row 9: Comentários
    setup_cell(9, 0, "Comentários:", bold=True)
    setup_cell(9, 1, "", span_to_col=9)
    table.rows[9].height = Mm(15)
    
    # Row 10: Data & Analista
    setup_cell(10, 0, "Data da Análise:", bold=True)
    setup_cell(10, 1, "", span_to_col=4)
    setup_cell(10, 5, "Analista GPA:", bold=True)
    setup_cell(10, 6, "", span_to_col=9)
    
    # Row 11: Header Quadro
    setup_cell(11, 0, "Análise do Quadro de Pessoal", bold=True, bg="D9D9D9", span_to_col=9)
    
    # Row 12: Quadro Aut, Atual, Vaga
    setup_cell(12, 0, "Quadro Autorizado:", bold=True)
    setup_cell(12, 1, "", span_to_col=2)
    setup_cell(12, 3, "Quadro Atual:", bold=True)
    setup_cell(12, 4, "", span_to_col=5)
    setup_cell(12, 6, "Nº da Vaga:", bold=True)
    setup_cell(12, 7, "", span_to_col=9)
    
    # Row 13: Comentários
    setup_cell(13, 0, "Comentários:", bold=True)
    setup_cell(13, 1, "", span_to_col=9)
    table.rows[13].height = Mm(15)
    
    # Row 14: Header Gerencia
    setup_cell(14, 0, "Preenchimento - Gerência Executiva Requisitante (Justificativa)", bold=True, bg="D9D9D9", span_to_col=9)
    
    # Row 15: Justificativa
    setup_cell(15, 0, "Justificativa:", bold=True)
    setup_cell(15, 1, "", span_to_col=9)
    table.rows[15].height = Mm(25)
    
    # Row 16: Candidato
    setup_cell(16, 0, "Candidato Aprovado:", bold=True)
    setup_cell(16, 1, "", span_to_col=3)
    setup_cell(16, 4, "Data:", bold=True)
    setup_cell(16, 5, "")
    setup_cell(16, 6, "Assinatura:", bold=True)
    setup_cell(16, 7, "", span_to_col=9)

    doc.add_paragraph() # Spacer
    
    # Flow Table
    flow_table = doc.add_table(rows=8, cols=2)
    flow_table.style = 'Table Grid'
    flow_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    flow_data = [
        ["JOVEM APRENDIZ", "TRAINEE/EFETIVO/GERENTE"],
        ["PARA APROVAÇÃO", "PARA APROVAÇÃO"],
        ["Gerente Executivo GP", "Gerente Imediato"],
        ["", "Gerente Mediato"],
        ["", "Diretor da Área"],
        ["", "Gerente Executivo GP"],
        ["", "Diretor DAC"],
        ["", "Diretor DGL"]
    ]
    
    for r, row in enumerate(flow_data):
        for c, text in enumerate(row):
            cell = flow_table.cell(r, c)
            cell.text = text
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(8)
            if r < 2:
                run.bold = True
                set_cell_background(cell, "D9D9D9")

    output_path = 'Solicitacao_Contratacao_Final.docx'
    doc.save(output_path)
    print(f"Documento final gerado: {output_path}")

if __name__ == "__main__":
    create_final_docx()
